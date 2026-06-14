#!/usr/bin/env python3
from pathlib import Path
import re
import shutil
import subprocess

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path("manuscript/TRUST-Swarm-HCC-manuscript-v1.md")
OUT_DIR = Path("submission")
DOCX = OUT_DIR / "TRUST-Swarm-HCC-final-submission.docx"
PDF = OUT_DIR / "TRUST-Swarm-HCC-final-submission.pdf"

if not SRC.exists():
    raise FileNotFoundError(SRC)

OUT_DIR.mkdir(parents=True, exist_ok=True)
raw = SRC.read_text()

def clean(s):
    s = s.replace("**", "")
    s = s.replace("`", "")
    s = s.replace("∈", "in")
    s = s.replace("Σ", "sum")
    s = s.replace("−", "-")
    return s.strip()

def set_font(run, size=10, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cols(section, num):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "360")

def is_table_line(line):
    return line.strip().startswith("|") and line.strip().endswith("|")

def is_sep(line):
    body = line.strip().strip("|").strip()
    return bool(body) and all(c in "-:| " for c in body)

def parse_table(lines):
    rows = []
    for line in lines:
        if is_sep(line):
            continue
        rows.append([clean(x) for x in line.strip().strip("|").split("|")])
    return rows

def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            txt = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_font(r, 7.5, bold=(i == 0))
            if i == 0:
                shade(cell, "D9EAF7")
    doc.add_paragraph()

def add_para(doc, text, size=9.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(clean(text))
    set_font(r, size)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(clean(text))
    set_font(r, 11 if level == 1 else 9.8, bold=True, italic=(level > 1))
    return p

def split_sections(md):
    matches = list(re.finditer(r"^##\s+(.+)$", md, re.MULTILINE))
    sections = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md)
        sections.append((m.group(1).strip(), md[start:end].strip()))
    return sections

title_match = re.search(r"^#\s+(.+)$", raw, re.MULTILINE)
title = title_match.group(1).strip() if title_match else "TRUST-Swarm"

sections = split_sections(raw)
section_map = {h: b for h, b in sections}

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(9.5)

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("High-Confidence Computing")
set_font(r, 13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Research article")
set_font(r, 9.5, italic=True)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
r = p.add_run(title)
set_font(r, 15, bold=True)

# Author block
p = doc.add_paragraph()
r = p.add_run("Sai Krishna Thota")
set_font(r, 10.5, bold=True)

p = doc.add_paragraph()
r = p.add_run("Independent Researcher, USA")
set_font(r, 9.5)

p = doc.add_paragraph()
r = p.add_run("Corresponding author: drsaikrishnathota@ieee.org")
set_font(r, 9)

# Article info + abstract table
front = doc.add_table(rows=1, cols=2)
front.alignment = WD_TABLE_ALIGNMENT.CENTER
front.style = "Table Grid"

left = front.cell(0, 0)
right = front.cell(0, 1)
shade(left, "F2F2F2")

left.text = ""
p = left.paragraphs[0]
r = p.add_run("ARTICLE INFO\n")
set_font(r, 9.5, bold=True)

r = p.add_run("\nArticle type:\nResearch article\n\nKeywords:\n")
set_font(r, 8.5)

keywords = section_map.get("Keywords", "").replace(";", "\n").splitlines()
for kw in keywords:
    kw = kw.strip()
    if kw:
        r = p.add_run(kw + "\n")
        set_font(r, 8.5)

right.text = ""
p = right.paragraphs[0]
r = p.add_run("ABSTRACT\n")
set_font(r, 9.5, bold=True)

abstract = section_map.get("Abstract", "")
abstract = re.sub(r"\n+", " ", abstract).strip()
r = p.add_run("\n" + clean(abstract))
set_font(r, 8.5)

# Body in two-column HCC-like layout
body_section = doc.add_section(WD_SECTION.CONTINUOUS)
set_cols(body_section, 2)

# Skip front matter sections already handled
skip = {"Abstract", "Keywords"}
for heading, body in sections:
    if heading in skip:
        continue

    add_heading(doc, heading, level=1)

    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if is_table_line(line):
            tlines = []
            while i < len(lines) and is_table_line(lines[i]):
                tlines.append(lines[i])
                i += 1
            add_table(doc, parse_table(tlines))
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:], level=2)
            i += 1
            continue

        if line.startswith("#### "):
            add_heading(doc, line[5:], level=2)
            i += 1
            continue

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean(line.strip()[2:]))
            set_font(r, 9)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(clean(re.sub(r"^\s*\d+\.\s+", "", line)))
            set_font(r, 9)
            i += 1
            continue

        if re.match(r"^\[\d+\]\s+", line):
            p = add_para(doc, line, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            i += 1
            continue

        add_para(doc, line)
        i += 1

doc.save(DOCX)
print(f"Saved DOCX: {DOCX}")

# Try PDF export
soffice_candidates = [
    shutil.which("soffice"),
    shutil.which("libreoffice"),
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
]
soffice = next((x for x in soffice_candidates if x and Path(x).exists()), None)

if soffice:
    subprocess.run([
        soffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(OUT_DIR),
        str(DOCX)
    ], check=False)
    generated = OUT_DIR / "TRUST-Swarm-HCC-final-submission.pdf"
    print(f"PDF export attempted: {generated}")
else:
    print("LibreOffice not found. DOCX created. Export PDF manually from Word if needed.")
